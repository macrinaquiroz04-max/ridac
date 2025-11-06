"""
📄 SERVICIO DE EXPORTACIÓN A WORD
Exporta análisis estructurado de documentos legales a formato Word
"""

from typing import Dict, List, Any, Optional
from datetime import datetime
import logging
import os
import tempfile
from collections import defaultdict

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger("word_export")

class WordExportService:
    """Servicio para exportar análisis a documentos Word"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx no está instalado. Ejecuta: pip install python-docx")
    
    def crear_reporte_analisis(self, 
                             resultado_analisis: Dict[str, Any], 
                             tomo_info: Dict[str, Any],
                             titulo: str = "Análisis de Carpeta de Investigación") -> str:
        """
        Crear reporte completo de análisis en formato Word
        
        Returns:
            str: Ruta del archivo Word generado
        """
        
        # Crear documento
        doc = Document()
        
        # Configurar estilos
        self._configurar_estilos(doc)
        
        # Agregar portada
        self._agregar_portada(doc, titulo, tomo_info)
        
        # Agregar resumen ejecutivo
        self._agregar_resumen_ejecutivo(doc, resultado_analisis)
        
        # Agregar secciones detalladas
        self._agregar_seccion_diligencias(doc, resultado_analisis.get('diligencias', []))
        self._agregar_seccion_fechas(doc, resultado_analisis.get('fechas', []))
        self._agregar_seccion_personas(doc, resultado_analisis.get('nombres', []))
        self._agregar_seccion_lugares(doc, resultado_analisis.get('lugares', []))
        self._agregar_seccion_direcciones(doc, resultado_analisis.get('direcciones', []))
        
        # Agregar alertas y estadísticas
        self._agregar_alertas(doc, resultado_analisis.get('alertas', []))
        self._agregar_estadisticas(doc, resultado_analisis)
        
        # Agregar metadatos de procesamiento
        self._agregar_metadatos(doc, resultado_analisis.get('metadatos', {}))
        
        # Guardar archivo temporal
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        logger.info(f"✅ Reporte Word generado: {temp_file.name}")
        return temp_file.name
    
    def _configurar_estilos(self, doc: Document):
        """Configurar estilos personalizados del documento"""
        
        # Estilo para títulos principales
        styles = doc.styles
        
        # Título principal
        if 'Título Principal' not in [s.name for s in styles]:
            titulo_style = styles.add_style('Título Principal', WD_STYLE_TYPE.PARAGRAPH)
            titulo_font = titulo_style.font
            titulo_font.name = 'Arial'
            titulo_font.size = Pt(18)
            titulo_font.bold = True
            titulo_font.color.rgb = None  # Color azul
            
        # Subtítulo
        if 'Subtítulo Custom' not in [s.name for s in styles]:
            subtitulo_style = styles.add_style('Subtítulo Custom', WD_STYLE_TYPE.PARAGRAPH)
            subtitulo_font = subtitulo_style.font
            subtitulo_font.name = 'Arial'
            subtitulo_font.size = Pt(14)
            subtitulo_font.bold = True
    
    def _agregar_portada(self, doc: Document, titulo: str, tomo_info: Dict):
        """Agregar portada del documento"""
        
        # Título principal
        titulo_p = doc.add_paragraph()
        titulo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo_run = titulo_p.add_run(titulo)
        titulo_run.font.size = Pt(20)
        titulo_run.font.bold = True
        
        doc.add_paragraph()  # Espacio
        
        # Información del tomo
        info_p = doc.add_paragraph()
        info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        info_text = f"""
        Carpeta de Investigación: {tomo_info.get('nombre_archivo', 'N/A')}
        Tomo ID: {tomo_info.get('id', 'N/A')}
        Fecha de Análisis: {datetime.now().strftime('%d de %B de %Y')}
        """
        
        info_run = info_p.add_run(info_text.strip())
        info_run.font.size = Pt(12)
        
        # Línea separadora
        doc.add_paragraph("=" * 60)
        doc.add_page_break()
    
    def _agregar_resumen_ejecutivo(self, doc: Document, resultado_analisis: Dict):
        """Agregar resumen ejecutivo"""
        
        doc.add_heading('RESUMEN EJECUTIVO', level=1)
        
        estadisticas = resultado_analisis.get('estadisticas', {})
        
        resumen_texto = f"""
        Este documento presenta el análisis automatizado de la carpeta de investigación, 
        extrayendo información relevante mediante técnicas de procesamiento de lenguaje natural 
        y reconocimiento de patrones.
        
        DATOS EXTRAÍDOS:
        • Fechas identificadas: {estadisticas.get('total_fechas', 0)}
        • Personas mencionadas: {estadisticas.get('total_nombres', 0)}
        • Direcciones encontradas: {estadisticas.get('total_direcciones', 0)}
        • Lugares referenciados: {estadisticas.get('total_lugares', 0)}
        • Diligencias procesadas: {estadisticas.get('total_diligencias', 0)}
        • Alertas generadas: {estadisticas.get('total_alertas', 0)}
        """
        
        doc.add_paragraph(resumen_texto.strip())
        doc.add_page_break()
    
    def _agregar_seccion_diligencias(self, doc: Document, diligencias: List[Dict]):
        """Agregar sección de diligencias cronológicamente ordenadas"""
        
        doc.add_heading('1. DILIGENCIAS PROCESALES', level=1)
        
        if not diligencias:
            doc.add_paragraph("No se encontraron diligencias en el documento.")
            return
        
        # Ordenar cronológicamente si tienen fecha
        diligencias_ordenadas = sorted(diligencias, 
                                     key=lambda x: x.get('fecha', ''), 
                                     reverse=False)
        
        counter = 1
        for diligencia in diligencias_ordenadas:
            doc.add_heading(f'1.{counter} {diligencia.get("tipo", "Diligencia")}', level=2)
            
            # Información básica
            tabla = doc.add_table(rows=4, cols=2)
            tabla.style = 'Table Grid'
            
            # Headers
            tabla.cell(0, 0).text = "Concepto"
            tabla.cell(0, 1).text = "Información"
            
            # Datos
            tabla.cell(1, 0).text = "Fecha"
            tabla.cell(1, 1).text = diligencia.get('fecha', 'No especificada')
            
            tabla.cell(2, 0).text = "Responsable"
            tabla.cell(2, 1).text = diligencia.get('responsable', 'No especificado')
            
            tabla.cell(3, 0).text = "Página"
            tabla.cell(3, 1).text = f"Página {diligencia.get('pagina', 'N/A')}"
            
            # Número de oficio si existe
            if diligencia.get('numero_oficio'):
                nueva_fila = tabla.add_row()
                nueva_fila.cells[0].text = "Número de Oficio"
                nueva_fila.cells[1].text = diligencia['numero_oficio']
            
            doc.add_paragraph()  # Espacio
            counter += 1
    
    def _agregar_seccion_fechas(self, doc: Document, fechas: List[Dict]):
        """Agregar sección de fechas cronológicamente"""
        
        doc.add_heading('2. CRONOLOGÍA DE EVENTOS', level=1)
        
        if not fechas:
            doc.add_paragraph("No se encontraron fechas relevantes.")
            return
        
        # Agrupar por tipo de fecha si es posible
        fechas_agrupadas = defaultdict(list)
        
        for fecha in fechas:
            texto_fecha = fecha.get('texto', '')
            pagina = fecha.get('pagina', 'N/A')
            
            # Intentar categorizar la fecha
            categoria = self._categorizar_fecha(texto_fecha)
            fechas_agrupadas[categoria].append({
                'texto': texto_fecha,
                'pagina': pagina
            })
        
        # Mostrar por categorías
        for categoria, lista_fechas in fechas_agrupadas.items():
            doc.add_heading(f'2.{len(fechas_agrupadas)} {categoria}', level=2)
            
            for fecha_info in lista_fechas:
                p = doc.add_paragraph()
                p.add_run(f"• {fecha_info['texto']}").bold = True
                p.add_run(f" (Página {fecha_info['pagina']})")
    
    def _agregar_seccion_personas(self, doc: Document, nombres: List[Dict]):
        """Agregar sección de personas con estadísticas"""
        
        doc.add_heading('3. PERSONAS IDENTIFICADAS', level=1)
        
        if not nombres:
            doc.add_paragraph("No se identificaron personas en el documento.")
            return
        
        # Contar menciones por persona
        contador_personas = defaultdict(list)
        
        for nombre in nombres:
            texto_nombre = nombre.get('texto', '')
            pagina = nombre.get('pagina', 'N/A')
            
            # Limpiar nombre para agrupar
            nombre_limpio = self._limpiar_nombre(texto_nombre)
            contador_personas[nombre_limpio].append(pagina)
        
        # Crear tabla de personas
        tabla = doc.add_table(rows=1, cols=3)
        tabla.style = 'Table Grid'
        
        # Headers
        hdr_cells = tabla.rows[0].cells
        hdr_cells[0].text = 'Nombre'
        hdr_cells[1].text = 'Menciones'
        hdr_cells[2].text = 'Páginas'
        
        # Ordenar por número de menciones
        personas_ordenadas = sorted(contador_personas.items(), 
                                  key=lambda x: len(x[1]), 
                                  reverse=True)
        
        for nombre, paginas in personas_ordenadas:
            row_cells = tabla.add_row().cells
            row_cells[0].text = nombre
            row_cells[1].text = str(len(paginas))
            row_cells[2].text = ', '.join(map(str, sorted(set(paginas))))
    
    def _agregar_seccion_lugares(self, doc: Document, lugares: List[Dict]):
        """Agregar sección de lugares"""
        
        doc.add_heading('4. LUGARES IDENTIFICADOS', level=1)
        
        if not lugares:
            doc.add_paragraph("No se identificaron lugares específicos.")
            return
        
        for lugar in lugares:
            p = doc.add_paragraph()
            p.add_run(f"• {lugar.get('texto', '')}").bold = True
            p.add_run(f" (Página {lugar.get('pagina', 'N/A')})")
    
    def _agregar_seccion_direcciones(self, doc: Document, direcciones: List[Dict]):
        """Agregar sección de direcciones"""
        
        doc.add_heading('5. DIRECCIONES IDENTIFICADAS', level=1)
        
        if not direcciones:
            doc.add_paragraph("No se identificaron direcciones específicas.")
            return
        
        for direccion in direcciones:
            p = doc.add_paragraph()
            p.add_run(f"• {direccion.get('texto', '')}").bold = True
            p.add_run(f" (Página {direccion.get('pagina', 'N/A')})")
    
    def _agregar_alertas(self, doc: Document, alertas: List[Dict]):
        """Agregar sección de alertas"""
        
        doc.add_heading('6. ALERTAS Y OBSERVACIONES', level=1)
        
        if not alertas:
            doc.add_paragraph("No se generaron alertas automáticas.")
            return
        
        for alerta in alertas:
            p = doc.add_paragraph()
            p.add_run(f"⚠️ {alerta.get('mensaje', '')}").bold = True
            if alerta.get('pagina'):
                p.add_run(f" (Página {alerta['pagina']})")
    
    def _agregar_estadisticas(self, doc: Document, resultado_analisis: Dict):
        """Agregar estadísticas generales"""
        
        doc.add_heading('7. ESTADÍSTICAS GENERALES', level=1)
        
        estadisticas = resultado_analisis.get('estadisticas', {})
        metadatos = resultado_analisis.get('metadatos', {})
        
        # Crear tabla de estadísticas
        tabla = doc.add_table(rows=1, cols=2)
        tabla.style = 'Table Grid'
        
        # Headers
        hdr_cells = tabla.rows[0].cells
        hdr_cells[0].text = 'Métrica'
        hdr_cells[1].text = 'Valor'
        
        # Datos
        metricas = [
            ('Total de fechas', estadisticas.get('total_fechas', 0)),
            ('Total de nombres', estadisticas.get('total_nombres', 0)),
            ('Total de direcciones', estadisticas.get('total_direcciones', 0)),
            ('Total de lugares', estadisticas.get('total_lugares', 0)),
            ('Total de diligencias', estadisticas.get('total_diligencias', 0)),
            ('Páginas procesadas', metadatos.get('numero_paginas', 0)),
            ('Tiempo de procesamiento', f"{metadatos.get('tiempo_procesamiento', 0):.2f} segundos"),
            ('Método de procesamiento', metadatos.get('metodo_procesamiento', 'N/A'))
        ]
        
        for metrica, valor in metricas:
            row_cells = tabla.add_row().cells
            row_cells[0].text = metrica
            row_cells[1].text = str(valor)
    
    def _agregar_metadatos(self, doc: Document, metadatos: Dict):
        """Agregar información técnica del procesamiento"""
        
        doc.add_heading('8. INFORMACIÓN TÉCNICA', level=1)
        
        info_tecnica = f"""
        Este reporte fue generado automáticamente mediante:
        
        • Motor de análisis: Sistema OCR Avanzado + NLP
        • Procesamiento: {metadatos.get('metodo_procesamiento', 'Chunks')}
        • Páginas analizadas: {metadatos.get('numero_paginas', 'N/A')}
        • Tiempo total: {metadatos.get('tiempo_procesamiento', 0):.2f} segundos
        • Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}
        
        NOTA: Esta herramienta utiliza inteligencia artificial para extraer información.
        Se recomienda verificar manualmente los datos críticos.
        """
        
        doc.add_paragraph(info_tecnica.strip())
    
    def _categorizar_fecha(self, texto_fecha: str) -> str:
        """Categorizar una fecha según su contexto"""
        texto_lower = texto_fecha.lower()
        
        if any(palabra in texto_lower for palabra in ['declaración', 'comparecencia', 'testimonio']):
            return 'Fechas de Diligencias'
        elif any(palabra in texto_lower for palabra in ['oficio', 'acuerdo']):
            return 'Fechas de Oficios'
        else:
            return 'Fechas Generales'
    
    def _limpiar_nombre(self, nombre: str) -> str:
        """Limpiar y normalizar nombre para agrupación"""
        # Remover títulos
        nombre = nombre.replace('C. ', '').replace('LIC. ', '').replace('DR. ', '')
        return nombre.strip().title()


# Instancia global del servicio
word_export_service = WordExportService() if DOCX_AVAILABLE else None
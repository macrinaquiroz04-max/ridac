# backend/app/services/export_service.py

from pathlib import Path
from typing import List
from datetime import datetime
from docx import Document
from openpyxl import Workbook
from sqlalchemy.orm import Session
from app.models.tomo import Tomo, ContenidoOCR
from app.models.carpeta import Carpeta
from app.config import settings
from app.utils.logger import logger

class ExportService:
    """Servicio de exportación a Word y Excel"""

    def __init__(self):
        self.export_path = Path(settings.EXPORT_PATH)
        self.export_path.mkdir(parents=True, exist_ok=True)

    def exportar_tomo_word(self, db: Session, tomo_id: int) -> str:
        """Exportar contenido OCR de un tomo a Word"""
        try:
            # Obtener tomo y contenido
            tomo = db.query(Tomo).filter(Tomo.id == tomo_id).first()
            if not tomo:
                raise Exception("Tomo no encontrado")

            contenidos = db.query(ContenidoOCR).filter(
                ContenidoOCR.tomo_id == tomo_id
            ).order_by(ContenidoOCR.pagina).all()

            # Crear documento Word
            doc = Document()

            # Título
            doc.add_heading(f'Documento OCR - {tomo.nombre_archivo}', 0)

            # Información del tomo
            doc.add_paragraph(f'Fecha de exportación: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
            doc.add_paragraph(f'Total de páginas: {tomo.total_paginas}')
            doc.add_paragraph(f'Motor OCR: {contenidos[0].motor_usado if contenidos else "N/A"}')
            doc.add_paragraph('')

            # Contenido por página
            for contenido in contenidos:
                doc.add_heading(f'Página {contenido.pagina}', level=2)
                doc.add_paragraph(f'Confianza: {contenido.confianza_porcentaje}%')
                doc.add_paragraph(contenido.texto_extraido or "(Sin texto extraído)")
                doc.add_page_break()

            # Guardar
            filename = f"tomo_{tomo_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = self.export_path / filename
            doc.save(str(filepath))

            logger.info(f"✓ Exportado a Word: {filepath}")
            return str(filepath)

        except Exception as e:
            logger.error(f"Error exportando a Word: {e}")
            raise

    def exportar_carpeta_excel(self, db: Session, carpeta_id: int) -> str:
        """Exportar índice de carpeta y tomos a Excel"""
        try:
            # Obtener carpeta y tomos
            carpeta = db.query(Carpeta).filter(Carpeta.id == carpeta_id).first()
            if not carpeta:
                raise Exception("Carpeta no encontrada")

            tomos = db.query(Tomo).filter(Tomo.carpeta_id == carpeta_id).all()

            # Crear libro Excel
            wb = Workbook()

            # Hoja 1: Información de la carpeta
            ws_info = wb.active
            ws_info.title = "Información"

            ws_info['A1'] = "Nombre de Carpeta"
            ws_info['B1'] = carpeta.nombre
            ws_info['A2'] = "Número de Expediente"
            ws_info['B2'] = carpeta.numero_expediente or "N/A"
            ws_info['A3'] = "Total de Tomos"
            ws_info['B3'] = len(tomos)
            ws_info['A4'] = "Fecha de Exportación"
            ws_info['B4'] = datetime.now().strftime("%Y-%m-%d %H:%M")

            # Hoja 2: Lista de tomos
            ws_tomos = wb.create_sheet("Tomos")
            headers = ["#", "Nombre Archivo", "Páginas", "Tamaño (MB)", "Estado OCR", "Progreso", "Fecha Subida"]
            ws_tomos.append(headers)

            for tomo in tomos:
                ws_tomos.append([
                    tomo.numero_tomo,
                    tomo.nombre_archivo,
                    tomo.total_paginas or 0,
                    float(tomo.tamaño_mb) if tomo.tamaño_mb else 0,
                    tomo.estado_ocr,
                    f"{tomo.progreso_ocr}%",
                    tomo.fecha_subida.strftime("%Y-%m-%d %H:%M") if tomo.fecha_subida else ""
                ])

            # Hoja 3: Contenido OCR (resumen)
            ws_ocr = wb.create_sheet("Resumen OCR")
            ws_ocr.append(["Tomo", "Página", "Confianza", "Preview Texto"])

            for tomo in tomos:
                contenidos = db.query(ContenidoOCR).filter(
                    ContenidoOCR.tomo_id == tomo.id
                ).limit(10).all()  # Solo primeras 10 páginas de cada tomo

                for contenido in contenidos:
                    preview = contenido.texto_extraido[:100] + "..." if contenido.texto_extraido and len(contenido.texto_extraido) > 100 else contenido.texto_extraido or ""
                    ws_ocr.append([
                        f"Tomo {tomo.numero_tomo}",
                        contenido.pagina,
                        f"{contenido.confianza_porcentaje}%",
                        preview
                    ])

            # Guardar
            filename = f"carpeta_{carpeta_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            filepath = self.export_path / filename
            wb.save(str(filepath))

            logger.info(f"✓ Exportado a Excel: {filepath}")
            return str(filepath)

        except Exception as e:
            logger.error(f"Error exportando a Excel: {e}")
            raise

    def exportar_busqueda_excel(self, resultados: List[dict]) -> str:
        """Exportar resultados de búsqueda a Excel"""
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Resultados de Búsqueda"

            # Headers
            headers = ["Carpeta", "Expediente", "Tomo", "Página", "Fragmento", "Confianza"]
            ws.append(headers)

            # Datos
            for resultado in resultados:
                ws.append([
                    resultado.get('carpeta_nombre', ''),
                    resultado.get('numero_expediente', ''),
                    resultado.get('tomo_numero', ''),
                    resultado.get('pagina', ''),
                    resultado.get('fragmento', ''),
                    resultado.get('confianza', '')
                ])

            # Guardar
            filename = f"busqueda_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            filepath = self.export_path / filename
            wb.save(str(filepath))

            logger.info(f"✓ Búsqueda exportada a Excel: {filepath}")
            return str(filepath)

        except Exception as e:
            logger.error(f"Error exportando búsqueda: {e}")
            raise
